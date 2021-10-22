"""

ЕСЛИ ВЫДЕЛЕННЫЙ ТЕКСТ УЖЕ БЫЛ СКРЫТЫМ, ОН ОСТАНЕТСЯ СКРЫТЫМ

не скрывает автоматически сгенерированное содержание и некоторые кривые списки
"""
import constants
import utils_py
from docx import Document, opc
from docx.enum.text import WD_COLOR_INDEX
from typing import List

tasks = {
    "1": "hide if colored in any color",
    "2": "hide if not colored in any color",
    "3": "hide if colored in chosen colors",
    "4": "hide if not colored in chosen colors",
    "5": "hide if italic",
    "6": "hide if not italic",
    "7": "hide if strike",
}


def get_colors() -> List[WD_COLOR_INDEX]:
    with open(f'{constants.utils_path}{constants.color_file}', "r", encoding="utf-8") as f:
        colors = f.read().split(",")
        # when file is empty, split returns empty string and list becomes of length of one
        if len(colors[0]) == 0:
            colors = []

    utils_py.log_in_file(f"GET_COLORS FUNCTION: colors: {colors} len(colors): {len(colors)}")
    wd_colors_list = [getattr(WD_COLOR_INDEX, color) for color in colors]
    return wd_colors_list


def get_pydocx_document(path_to_file_with_name, file_name):
    try:
        document = Document(path_to_file_with_name)
        return document
    except opc.exceptions.PackageNotFoundError:
        utils_py.log_in_file(
            f"HIDE FUNCTION: PackageNotFoundError {path_to_file_with_name} {file_name} - file is empty")

        Document().save(f'{constants.complete_path}{file_name}')
        return None


def change_font_properties(task, run, colors_list):
    if colors_list is None:
        colors_list = []
    if task == "1":
        if run.font.highlight_color:
            run.font.hidden = True
    elif task == "2":
        if run.font.highlight_color is None:
            run.font.hidden = True
    elif task == "3":
        if run.font.highlight_color in colors_list:
            run.font.hidden = True
    elif task == "4":
        if run.font.highlight_color not in colors_list:
            run.font.hidden = True
    elif task == "5":
        if run.font.italic:
            run.font.hidden = True
    elif task == "6":
        if run.font.italic is None:
            run.font.hidden = True
    elif task == "7":
        if run.font.strike:
            run.font.hidden = True


def iterate_document_part(document_part, task, colors_list):
    for t in document_part.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        change_font_properties(task, run, colors_list)

    for p in document_part.paragraphs:
        for run in p.runs:
            change_font_properties(task, run, colors_list)


def hide(path_to_file_with_name: str, file_name: str, task: str):
    """
         pydocx document is mutable, so we don't need to return it from functions
    """
    utils_py.log_in_file(f"HIDE FUNCTION: FILE PATH: {path_to_file_with_name} FILE NAME: {file_name}")
    utils_py.log_in_file(f"HIDE FUNCTION: TASK NUMBER: {task}, TASK DESCRIPTION: {tasks[task]}")

    document = get_pydocx_document(path_to_file_with_name, file_name)
    if document:

        if task in ["3", "4"]:
            colors_list = get_colors()
            if len(colors_list) == 0:
                utils_py.log_in_file(f"HIDE FUNCTION: no colors chosen")
                Document().save(f'{constants.complete_path}{file_name}')
                return
        else:
            colors_list = []

        iterate_document_part(document, task, colors_list)  # iterating through body

        # итерируемся по всем секциям в документе и обрабатываем все три вида верхних и нижних колонтитулов.
        # боковые колонтитулы тоже входят
        for section in document.sections:

            footers_and_headers = (section.footer, section.even_page_footer, section.first_page_footer,
                                   section.header, section.even_page_header, section.first_page_header)
            for footer_or_header in footers_and_headers:
                iterate_document_part(footer_or_header, task, colors_list)

        document.save(f'{constants.complete_path}{file_name}')
