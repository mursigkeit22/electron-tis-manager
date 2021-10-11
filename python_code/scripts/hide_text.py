from docx import oxml, Document, opc
from .. import constants, utils_py


def hide_text(path_to_file_with_name, file_name):
    utils_py.log_in_file(f"HIDE_TEXT: {path_to_file_with_name} {file_name}")

    try:
        document = Document(path_to_file_with_name)
    except opc.exceptions.PackageNotFoundError:
        utils_py.log_in_file(f"HIDE_TEXT: PackageNotFoundError {path_to_file_with_name} {file_name}")

        Document().save(f'{constants.complete_path}{file_name}')
        return

    # текст документа
    elements_list = document._element.xpath('//w:r')

    WPML_URI = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    tag_rPr = WPML_URI + 'rPr'
    tag_highlight = WPML_URI + 'highlight'
    tag_t = WPML_URI + 't'

    def hide_not_highlighted_text(elements):
        """
        не используем стандартные методы библиотеки docx,
        так как с их помощью нельзя скрыть автоматически сгенерированное содержание.

        """
        utils_py.log_in_file(f"HIDE_TEXT: hide_not_highlighted_text")
        for element in elements:
            try:
                rpr_list = element.findall(tag_rPr)
                if len(rpr_list) == 0:
                    tag_text = element.find(tag_t)
                    rpr = oxml.shared.OxmlElement('w:rPr')  # создаем тег rPr
                    rpr.append(oxml.shared.OxmlElement('w:vanish'))  # добавляем в него тег скрытого текста
                    tag_text.addprevious(rpr)  # добавляем новый тег выше тега с текстом

                for rPr in rpr_list:  # если уже есть тег rPr
                    high = rPr.findall(tag_highlight)  # проверяем, есть ли тег, отвечающий за подсветку
                    if len(high) == 0:
                        rPr.append(
                            oxml.shared.OxmlElement('w:vanish'))  # если нет, добавляем в него тег скрытого текста

            except AttributeError:  # if there is no text
                pass

    def iterate_footer_header(footer_or_header):
        """
        используем методы  docx библиотеки, чтобы проверить весь текст в колонтитулах.
        работаем с таблицами и с параграфами, которые могут быть в колонтитулах.
        """
        utils_py.log_in_file(f"HIDE_TEXT: iterate_footer_header {footer_or_header}")

        for t in footer_or_header.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if run.font.highlight_color is None:
                                run.font.hidden = True
        for p in footer_or_header.paragraphs:
            for run in p.runs:
                if run.font.highlight_color is None:
                    run.font.hidden = True

    # скрываем невыделенный текст везде кроме колонтитулов
    hide_not_highlighted_text(elements_list)

    # итерируемся по всем секциям в документе и обрабатываем все три вида верхних и нижних колонтитулов.
    # боковые колонтитулы тоже входят

    for section in document.sections:
        utils_py.log_in_file(f"HIDE_TEXT: iterating through sections")

        footers_and_headers = (section.footer, section.even_page_footer, section.first_page_footer,
                               section.header, section.even_page_header, section.first_page_header)
        for footer_or_header in footers_and_headers:
            iterate_footer_header(footer_or_header)

    document.save(f'{constants.complete_path}{file_name}')
